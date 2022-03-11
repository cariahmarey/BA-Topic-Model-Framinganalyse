import os
from docx import Document
import csv
import pandas as pd


# if __name__ == '__main__':
#    docxpath = "Docx\DocxEnglisch"
#    doc_list = os.listdir(docxpath)
#
#    for doc in doc_list:
#        path = os.path.join(docxpath, doc)
#        document = Document(path)
#        i = 0
#        for p in document.paragraphs[1:]:
#            if not len(l.text) == 0 and 'S.' in p.text and len(p.text) > 40:
#                i += 1
#                print(f'{doc} {p.text}')
                
###########################
if __name__ == '__main__':

    # Pfad
    docxpath = "Docx\DocXEnglisch"
    doc_list = os.listdir(docxpath)

    output_list = []

    for doc in doc_list:
        path = os.path.join(docxpath, doc)
        document = Document(path)
        for p in document.paragraphs[1:]:
            if not len(p.text.strip()) == 0 and 'S.' not in p.text: #search for 'S.' and linebreak
                text = p.text.replace("\n", " ")
                output_list.append([doc, text])

    # create dataframe
    output_df = pd.DataFrame(output_list, columns=['File', 'Text'])
    # create csv from dataframe
    output_df.to_csv('corpus_englisch.csv',
                     index=False,
                     sep=';',
                     encoding='utf-8-sig')


